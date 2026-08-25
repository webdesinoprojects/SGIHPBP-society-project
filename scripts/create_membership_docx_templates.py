"""Create the SGIHPBP receipt/certificate DOCX templates and local QA samples."""

from __future__ import annotations

import argparse
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "frontend" / "src" / "assets"
TEMPLATE_DIR = ASSETS / "user-docs"
CERTIFICATE_PATH = TEMPLATE_DIR / "Certificate_Format.docx"
RECEIPT_PATH = TEMPLATE_DIR / "Receipt Format.docx"
LOGO_PATH = ASSETS / "Logo_SGIHPBPS.png"
TEAL = "0B3B46"


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.1)
    section.footer_distance = Inches(0.1)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)


def set_cell_border(cell, **edges) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, attributes in edges.items():
        tag = f"w:{edge_name}"
        edge = borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            borders.append(edge)
        for key, value in attributes.items():
            edge.set(qn(f"w:{key}"), str(value))


def remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                bottom={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                insideH={"val": "nil"},
                insideV={"val": "nil"},
            )


def set_cell_margins(cell, top=0, start=50, bottom=0, end=50) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_run(paragraph, text, *, size=11, bold=False, italic=False, underline=False, font="Arial", color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_brand_header(container, compact=False) -> None:
    table = container.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.45 if compact else 1.55)
    table.columns[1].width = Inches(8.7)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    set_cell_margins(left, start=0, end=80)
    set_cell_margins(right, start=80, end=0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    logo_p = left.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.25 if compact else 1.35))

    title_p = right.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(3)
    add_run(
        title_p,
        "Society of Gastrointestinal &\nHepatopancreatobiliary Pathologists, of India",
        size=21 if compact else 23,
        bold=True,
        font="Arial Narrow",
    )
    registration_p = right.add_paragraph()
    registration_p.paragraph_format.space_after = Pt(0)
    add_run(
        registration_p,
        "Registered under the societies registration act 1860 (Delhi) No. S/4105/SDM/NW/2024",
        size=8.5,
        bold=True,
    )


def extract_legacy_signatures() -> dict[str, Path]:
    supplied = {
        "president.png": TEMPLATE_DIR / "signatures" / "puja-sakhuja.png",
        "vice-president.png": TEMPLATE_DIR / "signatures" / "anjali-amarapurkar.png",
        "secretary.png": TEMPLATE_DIR / "signatures" / "prasenjit-das.png",
        "treasurer.jpeg": TEMPLATE_DIR / "signatures" / "arvind-ahuja.png",
    }
    if all(path.exists() for path in supplied.values()):
        return supplied

    work = Path(tempfile.mkdtemp(prefix="sgihpbp-signatures-"))
    extracted: dict[str, Path] = {}
    sources = [
        (CERTIFICATE_PATH, "word/media/image2.png", "president.png"),
        (CERTIFICATE_PATH, "word/media/image4.jpeg", "treasurer.jpeg"),
        (CERTIFICATE_PATH, "word/media/image3.png", "secretary.png"),
    ]
    for archive, entry, name in sources:
        if not archive.exists():
            continue
        try:
            with zipfile.ZipFile(archive) as package:
                target = work / name
                target.write_bytes(package.read(entry))
                extracted[name] = target
        except (KeyError, zipfile.BadZipFile):
            continue
    return extracted


def add_signature(cell, image_path, name, role, *, text_signature=None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    set_cell_margins(cell, start=35, end=35)
    signature_p = cell.paragraphs[0]
    signature_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature_p.paragraph_format.space_after = Pt(1)
    if image_path and image_path.exists():
        signature_p.add_run().add_picture(str(image_path), width=Inches(0.8))
    else:
        add_run(signature_p, text_signature or name, size=10, italic=True, font="Segoe Script")
    name_p = cell.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(0)
    add_run(name_p, name, size=9.5, font="Times New Roman")
    role_p = cell.add_paragraph()
    role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    role_p.paragraph_format.space_after = Pt(0)
    add_run(role_p, role, size=9.5, font="Times New Roman")


def create_certificate(signatures: dict[str, Path]) -> None:
    document = Document()
    configure_page(document)
    add_brand_header(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(18)
    add_run(title, "{{MEMBERSHIP_TITLE}}", size=20, font="Book Antiqua")

    member = document.add_paragraph()
    member.alignment = WD_ALIGN_PARAGRAPH.CENTER
    member.paragraph_format.space_after = Pt(26)
    add_run(member, "{{MEMBER_LABEL}}", size=19, font="Book Antiqua")

    line = document.add_paragraph()
    line.paragraph_format.space_after = Pt(6)
    p_pr = line._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)

    table = document.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)
    names = [
        (signatures.get("president.png"), "Professor Puja Sakhuja", "President", "Puja Sakhuja"),
        (signatures.get("vice-president.png"), "Prof. Anjali Amarapurkar", "Vice- President", "A. D. Amarapurkar"),
        (signatures.get("secretary.png"), "Professor Prasenjit Das", "Secretary General", None),
        (signatures.get("treasurer.jpeg"), "Professor Arvind Ahuja", "Treasurer", None),
    ]
    for cell, (image, name, role, text_signature) in zip(table.rows[0].cells, names):
        add_signature(cell, image, name, role, text_signature=text_signature)

    document.save(CERTIFICATE_PATH)


def create_receipt(signatures: dict[str, Path]) -> None:
    document = Document()
    configure_page(document)
    outer = document.add_table(rows=1, cols=1)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    cell = outer.cell(0, 0)
    set_cell_margins(cell, top=70, start=130, bottom=70, end=130)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "18", "color": TEAL},
        bottom={"val": "single", "sz": "18", "color": TEAL},
        left={"val": "single", "sz": "18", "color": TEAL},
        right={"val": "single", "sz": "18", "color": TEAL},
    )

    bill = cell.paragraphs[0]
    bill.paragraph_format.space_after = Pt(3)
    add_run(bill, "Bill No: {{BILL_NUMBER}}", size=9)
    add_brand_header(cell, compact=True)
    pan = cell.add_paragraph()
    pan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pan.paragraph_format.space_after = Pt(13)
    add_run(pan, "PAN ACMAS8047C, NPODARPAN ID DL/2025/0619963", size=8.5, bold=True)

    body = cell.add_paragraph()
    body.paragraph_format.left_indent = Inches(0.42)
    body.paragraph_format.right_indent = Inches(0.35)
    body.paragraph_format.space_after = Pt(7)
    body.paragraph_format.line_spacing = 1.08
    add_run(body, "This is to certify that a total sum of ", size=12)
    add_run(body, "{{AMOUNT}}", size=12, underline=True)
    add_run(body, " has been received from ", size=12)
    add_run(body, "{{APPLICANT_NAME}}", size=12, underline=True)
    add_run(body, ", transaction ID {{TRANSACTION_DETAILS}}, on account of his/her inclusion as a {{MEMBER_TYPE}} of the Society.", size=12)

    membership = cell.add_paragraph()
    membership.paragraph_format.left_indent = Inches(0.42)
    membership.paragraph_format.space_after = Pt(13)
    add_run(membership, "The membership number of {{MEMBERSHIP_NUMBER}} has been allotted to him/her.", size=12)

    auth = cell.add_paragraph()
    auth.paragraph_format.left_indent = Inches(0.42)
    auth.paragraph_format.space_after = Pt(2)
    add_run(auth, "Authorized Signatories", size=11)

    signatures_table = cell.add_table(rows=1, cols=2)
    signatures_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    signatures_table.autofit = False
    remove_table_borders(signatures_table)
    signatures_table.columns[0].width = Inches(3.2)
    signatures_table.columns[1].width = Inches(3.2)
    add_signature(signatures_table.cell(0, 0), signatures.get("treasurer.jpeg"), "Treasurer, SGIHPBPs", "")
    add_signature(signatures_table.cell(0, 1), signatures.get("secretary.png"), "Secretary General, SGIHPBPs", "")

    footer = cell.add_table(rows=1, cols=2)
    footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(footer)
    left = footer.cell(0, 0).paragraphs[0]
    right = footer.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(left, "e-PAN No: ACMAS8047C", size=10, font="Times New Roman")
    add_run(right, "Dated: {{APPROVED_DATE}}", size=10, font="Times New Roman")

    document.save(RECEIPT_PATH)


def replace_tokens(source: Path, target: Path, values: dict[str, str]) -> None:
    with zipfile.ZipFile(source) as package:
        with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as output:
            for item in package.infolist():
                data = package.read(item.filename)
                if item.filename == "word/document.xml":
                    xml = data.decode("utf-8")
                    for token, value in values.items():
                        xml = xml.replace(token, value)
                    data = xml.encode("utf-8")
                output.writestr(item, data)


def create_samples(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    common = {
        "{{APPROVED_DATE}}": "25/08/2026",
        "{{APPLICANT_NAME}}": "DR B PARTHEEBAN",
        "{{TRANSACTION_DETAILS}}": "525149781378 Google Pay",
        "{{AMOUNT}}": "INR 10,000/-",
        "{{MEMBER_TYPE}}": "Life Member",
        "{{BILL_NUMBER}}": "00038/LM/2026",
        "{{MEMBERSHIP_NUMBER}}": "SGIHPBP/00008/GM/2025",
        "{{MEMBERSHIP_TITLE}}": "LIFE MEMBERSHIP",
        "{{MEMBER_LABEL}}": "DR B PARTHEEBAN (SGIHPBP/00008/GM/2025)",
    }
    replace_tokens(RECEIPT_PATH, output_dir / "sample-life-receipt.docx", common)
    replace_tokens(CERTIFICATE_PATH, output_dir / "sample-life-certificate.docx", common)

    adhoc = {
        **common,
        "{{APPLICANT_NAME}}": "DR ROMA PAUDAL",
        "{{TRANSACTION_DETAILS}}": "600796005675 Date: 07/01/2026",
        "{{AMOUNT}}": "INR 2,500/-",
        "{{MEMBER_TYPE}}": "Ad Hoc Member",
        "{{BILL_NUMBER}}": "00160/AH/2026",
        "{{MEMBERSHIP_NUMBER}}": "SGIHPBP/00034AdM/2025",
        "{{MEMBERSHIP_TITLE}}": "AD HOC MEMBERSHIP (FOR 3 YEARS)",
        "{{MEMBER_LABEL}}": "DR ROMA PAUDAL (SGIHPBP/00034AdM/2025)",
    }
    replace_tokens(RECEIPT_PATH, output_dir / "sample-ad-hoc-receipt.docx", adhoc)
    replace_tokens(CERTIFICATE_PATH, output_dir / "sample-ad-hoc-certificate.docx", adhoc)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--samples-dir", type=Path)
    args = parser.parse_args()

    signatures = extract_legacy_signatures()
    create_certificate(signatures)
    create_receipt(signatures)
    if args.samples_dir:
        create_samples(args.samples_dir)


if __name__ == "__main__":
    main()
